#!/usr/bin/env python
# coding: utf-8

# In[ ]:


### Imports ###
import urllib.request, urllib.parse, urllib.error
from bs4 import BeautifulSoup
import ssl
import docx
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import csv

### Ignoring Errors ###
ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE

### Enter Reference No. ###
print("**********Are you going to enter the Reference No. 'Manually' or going to use 'CSV file'**********")
print('Enter "M" for manual')
print('Enter "C" for CSV file')
e_m_mode = 0
R = list()
while e_m_mode == 0:
    entry_mode = input('Entry Mode ? ')
    if entry_mode.upper() == "M" or entry_mode.upper() == "C":
        if entry_mode.upper() == "M":
            check_R_m = 0
            while check_R_m == 0:
                rn = input('Enter 15 Digit Reference No. ')
                if len(rn) == 15:
                    R.append(rn)
                    check_R_m = 1
                else:
                    continue        
        if entry_mode.upper() == "C":            
            check_R_c = 0
            while check_R_c == 0:
                csv_file_name = input("Enter CSV file name: ")
                try:
                    with open(f"{csv_file_name}.csv") as f:
                        contents_of_f = csv.reader(f)
                        for each_line in contents_of_f:
                            R += each_line
                    if len(R) >= 1:
                        check_R_c = 1
                    else:
                        print("Entered file contains no entry, Kindly Enter valid file name.")
                except:
                    check_R_c = 0            
        e_m_mode = 1
    else:
        print("Kindly enter 'M' or 'C'")
R_ok = list()
R_error = list()
for elem in R:
    if len(elem) == 15:
        R_ok.append(elem)
    else:
        R_error.append(elem)

executioncount = 0
filecount = 0
reference_number_causing_errors = list()
reference_number_giving_no_result = list()
for elem_R in R_ok:
    executioncount += 1
### Process Reference No. ###
    reference_no = elem_R
    BatchNo = reference_no[0:2]
    SubDiv = reference_no[2:7]
    RefNo = reference_no[7:14]
    RU = reference_no[14]
    RU = RU.upper()
    try:
        ### Read URL data ###
        url = f"http://lesco.gov.pk/Modules/CustomerBill/BillPrintMDI.asp?nBatchNo={BatchNo}&nSubDiv={SubDiv}&nRefNo={RefNo}&strRU={RU}"
        html = urllib.request.urlopen(url, context=ctx).read()
        soup = BeautifulSoup(html, 'html.parser')

        ### Reading Reference No. ###
        tags = soup('div')
        n = 0
        m = 0
        reference_no1 = 0
        for tag in tags:
            a = tag.contents[0]
            try:
                a = a.lstrip()
                a = a.rstrip()
                if len(a) <= 0:
                    continue
                else:
                    n += 1
                    if a == "OLD A/C No.":
                        b = n+1
                    elif n == b:
                        reference_no1 = a
            except:
                continue

        ### Reading Name, Address and Current MDI ###
        tags = soup('td')
        n = 0
        name = ""
        address = ""
        for tag in tags:
            a = tag.contents[0]
            try:
                a = a.lstrip()
                a = a.rstrip()
                if len(a) <= 0:
                    continue
                else:
                    n = n+1
                    if a == "NAME & ADDRESS":
                        b = n+1
                    if n == b:
                        name = a
                    if n == b+1:
                        address = a
            except:
                continue
        print("")
        print("")
        print("**********************************************************************************************")
        print("***                                                                                        ***")
        print("***   Kindly verify the data")
        print(f"***   Reference No. : {reference_no1}")
        print(f"***   Name : {name}")
        print(f"***   Address : {address}")
            
        ###Open Docx.###
        document = docx.Document()

        ###Generate File Name###
        filename = f"MDI Notice {BatchNo}.{SubDiv}.{RefNo}"

        ### Set margins ###
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1.78)
            section.right_margin = Cm(1.78)

        ### Set Styles ###
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        ### ADD header ###
        document.add_picture("11221-head.jpg", width = docx.shared.Cm(18), height = docx.shared.Cm(2.5)) 

        ### Memo No. and Date ###
        paragraph = document.add_paragraph("Memo No. __________                                                                                              Dated ____/____/2021.")
        paragraph = document.add_paragraph("")

        ### Recievers ###
        paragraph = document.add_paragraph("To,")
        paragraph = document.add_paragraph(f"{name}")
        paragraph.paragraph_format.left_indent = Inches(1)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph = document.add_paragraph(f"{address}")
        paragraph.paragraph_format.left_indent = Inches(1)
        paragraph = document.add_paragraph("")

        ### Subject ###
        subject = f"NOTICE  UNDER  SECTION  (26-A) & SECTION  (24)  OF  ELECTRICITY  ACT-"
        subject3 = f"1910 AS AMENDED AGAINST REFERENCE NO. {SubDiv}.{RefNo}/{BatchNo} IN THE NAME OF {name} R/O {address}."
        paragraph = document.add_paragraph(f"Subject: \t")
        runner = paragraph.add_run(subject)
        runner.bold = True
        runner.underline = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(1)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        runner = paragraph.add_run(subject3)
        runner.bold = True
        runner.underline = True
        paragraph = document.add_paragraph("")

        ### Body ##
        paragraph = document.add_paragraph(f"\tIt is informed that your premises were checked ")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #
        paragraph = document.add_paragraph("\tYou are therefore advised to coordinate with this office for assessment of the loss sustained to the company due to said discrepancy under section-26 of Elect: Act-1910 as amended within 07-days.")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #
        paragraph = document.add_paragraph("\tIn case of non-representation from your end within stipulated period, Ex-part action will be taken and also further action under section-24 of Elec: Act 1910 will be initiated for disconnection of supply at premises.")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        ### Footer ###
        paragraph = document.add_paragraph("")
        paragraph = document.add_paragraph("")
        off = "Addl. EXECUTIVE ENGINEER"
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        runner = paragraph.add_run(off)
        runner.bold = True
        paragraph.paragraph_format.left_indent = Inches(4.4)
        sub = "ALI RAZA ABAD SUB DIVISION LESCO"
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = paragraph.add_run(sub)
        runner.bold = True


        ### CC##
        paragraph = document.add_paragraph("")
        paragraph = document.add_paragraph("")
        paragraph = document.add_paragraph("C.C")
        paragraph = document.add_paragraph("\tMaster file.")

        ### Saving File ###
        if reference_no1 != 0 :
            document.save(f"{filename}.docx")
            print(f"***   {filename}.docx Created")
            filecount +=1
        else :
            reference_number_giving_no_result.append(reference_no)
        print("***                                                                                        ***")
        print("**********************************************************************************************")
        print("")
        print("")
    except:
        reference_number_causing_errors.append(reference_no)

        
print("----------------------------------------------------------------------------------------------")
print("----------------------------------------------------------------------------------------------")
print("----------------------------------------------------------------------------------------------")

print(" ")
print(f"Total reference nos. entered or in file : {len(R)}")
print(f"Total Number of records processed : {executioncount}")      
print(f"Number of files created : {filecount}") 
if len(R_error) > 0:
    n = 1
    print(" ")
    print("Following Reference Nos. in the CSV are invalid kindly re-check them manually")
    for elem in R_error:
        print (f"{n} --- {elem}")
        n += 1
if len(reference_number_causing_errors) > 0:
    n = 1
    print(" ")
    print("Following Reference Nos. are generating no result kindly re-check them manually")
    for elem in reference_number_causing_errors:
        print (f"{n} --- {elem}")
        n += 1
if len(reference_number_giving_no_result) > 0:
    n = 1
    print(" ")
    print("Following Reference Nos. are generating no result kindly re-check them manually")
    for elem in reference_number_giving_no_result:
        print (f"{n} --- {elem}")
        n += 1

